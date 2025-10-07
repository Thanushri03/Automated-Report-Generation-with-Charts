

# app----final.py
import os
import io
import json
import pandas as pd
import matplotlib.pyplot as plt
from typing import List, Tuple
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
import streamlit as st

# Google Gemini (google-genai)
from google import genai
from google.genai import types
import re

def clean_text(text: str) -> str:
    """Remove NULL bytes and non-XML-compatible control characters."""
    # Remove null bytes and all control chars except common whitespace
    return re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", "", text)


# ===============================
# Streamlit Page Config
# ===============================
st.set_page_config(
    page_title="Report Generator",
    page_icon="📑",
    layout="wide"
)

st.title("📑 Report Generator (PDF/DOCX/CSV → Gemini → DOCX)")

st.markdown(
    "Upload **multiple** internal files (PDF/DOCX/CSV). "
    "The app will extract text, chunk it, send context to **Gemini** for a structured report, "
    "render any suggested charts, and provide a **downloadable DOCX**."
)

# ===============================
# Sidebar Controls
# ===============================
with st.sidebar:
    st.header("⚙️ Settings")

    # API Key handling
    default_key = os.environ.get("GEMINI_API_KEY", "")
    api_key = st.text_input("Gemini API Key", value=default_key, type="password", help="You can also set GEMINI_API_KEY env var.")
    model_name = st.text_input("Model Name", value="gemini-2.5-flash", help="e.g., gemini-2.5-flash")

    # Chunking
    st.subheader("Chunking")
    chunk_size = st.number_input("Chunk size", min_value=200, max_value=2000, value=400, step=50)
    overlap = st.number_input("Overlap", min_value=0, max_value=500, value=80, step=10)

    # Optional extra external context
    st.subheader("External Web Inputs")
    external_text = st.text_area(
        "Add optional external notes/market data (will be appended to context)",
        value=(
            "Recent market analysis indicates a 15% increase in commodity prices globally, expected "
            "to impact future costs, particularly in manufacturing-heavy regions like Region A.\n"
            "Competitor 'X' released a highly-rated alternative to our Product Alpha.\n"
            "Recommendation from analysts: allocate 60% of next marketing budget to Region C to defend share."
        ),
        height=140
    )

# ===============================
# Helpers: Chunkers
# ===============================
def pdf_to_chunks_bytes(data: bytes, chunk_size: int = 400, overlap: int = 80, file_label: str = "PDF") -> List[Tuple[str, str]]:
    chunks = []
    reader = PdfReader(io.BytesIO(data))
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = clean_text(text)  # ✅ clean here
        start = 0
        while start < len(text):
            chunk = text[start:start + chunk_size].strip()
            if chunk:
                chunks.append((chunk, f"{file_label} p.{i+1}"))
            start += max(1, chunk_size - overlap)
    return chunks


def docx_to_chunks_bytes(data: bytes, chunk_size: int = 400, overlap: int = 80, file_label: str = "DOCX") -> List[Tuple[str, str]]:
    """
    Extract text from a DOCX bytes and split into chunks.
    Returns list of (chunk_text, location_label).
    """
    doc = Document(io.BytesIO(data))
    text = "\n".join([para.text for para in doc.paragraphs])
    text = clean_text(text)
    chunks: List[Tuple[str, str]] = []
    start = 0
    while start < len(text):
        chunk = text[start:start + chunk_size].strip()
        if chunk:
            loc = f"{file_label}"
            chunks.append((chunk, loc))
        start += max(1, chunk_size - overlap)
    return chunks


def csv_to_chunks_bytes(data: bytes, chunk_size: int = 400, overlap: int = 80, file_label: str = "CSV") -> List[Tuple[str, str]]:
    """
    Read a CSV from bytes and convert rows to text chunks.
    Returns list of (chunk_text, location_label).
    """
    df = pd.read_csv(io.BytesIO(data))
    text = df.to_string(index=False)
    chunks: List[Tuple[str, str]] = []
    start = 0
    while start < len(text):
        chunk = text[start:start + chunk_size].strip()
        if chunk:
            loc = f"{file_label}"
            chunks.append((chunk, loc))
        start += max(1, chunk_size - overlap)
    return chunks


def load_multiple_files(files, chunk_size: int, overlap: int) -> str:
    """
    Load multiple uploaded files, extract and chunk text, and combine to one context string.
    """
    all_chunks: List[Tuple[str, str]] = []

    for f in files:
        name = f.name
        ext = os.path.splitext(name)[1].lower()
        data = f.read()
        label = f"{name}"

        if ext == ".pdf":
            chunks = pdf_to_chunks_bytes(data, chunk_size=chunk_size, overlap=overlap, file_label=label)
        elif ext == ".docx":
            chunks = docx_to_chunks_bytes(data, chunk_size=chunk_size, overlap=overlap, file_label=label)
        elif ext == ".csv":
            chunks = csv_to_chunks_bytes(data, chunk_size=chunk_size, overlap=overlap, file_label=label)
        else:
            st.warning(f"Unsupported file type skipped: {name}")
            continue

        all_chunks.extend(chunks)

    # Combine
    combined_text = "\n".join([f"[{loc}] {c}" for c, loc in all_chunks])
    return combined_text


# ===============================
# JSON Schema for Gemini
# ===============================
REPORT_SCHEMA = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "executive_summary": {"type": "string"},
        "key_insights": {"type": "array", "items": {"type": "string"}},
        "data_for_charts": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "chart_title": {"type": "string"},
                    "x_labels": {"type": "array", "items": {"type": "string"}},
                    "y_data": {"type": "array", "items": {"type": "number"}},
                    "chart_type_suggestion": {"type": "string", "enum": ["bar", "pie", "line", "table"]}
                },
                "required": ["chart_title", "x_labels", "y_data", "chart_type_suggestion"]
            }
        },
        "recommendations": {"type": "array", "items": {"type": "string"}},
        "citations": {"type": "array", "items": {"type": "string"}}
    },
    "required": [
        "title",
        "executive_summary",
        "key_insights",
        "data_for_charts",
        "recommendations",
        "citations"
    ]
}


# ===============================
# Gemini Call
# ===============================
def build_gemini_client(api_key: str):
    if not api_key:
        raise ValueError("Missing Gemini API key. Provide it in the sidebar or set GEMINI_API_KEY.")
    # prefer explicit api_key over env var to avoid confusion
    return genai.Client(api_key=api_key)


def generate_structured_report(client, model_name: str, context: str) -> dict:
    prompt = (
        "Analyze the following data context. Extract key findings, numerical data, and synthesize "
        "a comprehensive structured report following the provided JSON schema. "
        "Prefer creating one or two charts if numeric data is present. "
        "Return only valid JSON per the schema.\n\n"
        f"CONTEXT:\n{context}"
    )
    config = types.GenerateContentConfig(
        response_mime_type="application/json",
        response_schema=REPORT_SCHEMA
    )
    response = client.models.generate_content(
        model=model_name,
        contents=[prompt],
        config=config
    )
    return json.loads(response.text)


# ===============================
# Charts
# ===============================
def generate_chart(chart_data: dict) -> io.BytesIO | None:
    """Generate Matplotlib chart → PNG bytes buffer."""
    title = chart_data["chart_title"]
    x_labels = chart_data["x_labels"]
    y_data = chart_data["y_data"]
    chart_type = chart_data["chart_type_suggestion"].lower()

    fig, ax = plt.subplots(figsize=(7, 4))
    try:
        if chart_type == "bar":
            ax.bar(x_labels, y_data)
            ax.set_ylabel("Value")
        elif chart_type == "pie":
            ax.pie(y_data, labels=x_labels, autopct="%1.1f%%", startangle=90)
            ax.axis("equal")
        elif chart_type == "line":
            ax.plot(x_labels, y_data, marker="o")
            ax.set_ylabel("Trend")
        else:
            # fallback to bar
            ax.bar(x_labels, y_data)
            ax.set_ylabel("Value")

        ax.set_title(title)
        plt.xticks(rotation=20, ha="right")
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format="png")
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as e:
        plt.close(fig)
        st.error(f"Chart generation failed for '{title}': {e}")
        return None


# ===============================
# DOCX Assembly → Bytes
# ===============================
def create_docx_report_bytes(report_data: dict, chart_buffers: list) -> bytes:
    """Assemble the DOCX in-memory and return bytes."""
    document = Document()
    document.add_heading(report_data.get("title", "Automated Report"), 0)
    document.add_paragraph(f"Report Date: {pd.Timestamp.now().strftime('%B %d, %Y')}")
    document.add_page_break()

    # Executive Summary
    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(report_data.get("executive_summary", "No summary provided."))

    # Key Insights
    document.add_heading("Key Insights", level=1)
    for insight in report_data.get("key_insights", []):
        document.add_paragraph(insight, style="List Bullet")

    # Visualizations
    document.add_heading("Data Visualizations", level=1)
    for i, chart_info in enumerate(report_data.get("data_for_charts", [])):
        if chart_info["chart_type_suggestion"].lower() == "table":
            document.add_paragraph(f"Table {i+1}: {chart_info['chart_title']}")
            table = document.add_table(rows=len(chart_info["x_labels"]) + 1, cols=2)
            table.style = "Light Shading"
            hdr = table.rows[0].cells
            hdr[0].text = "Metric"
            hdr[1].text = "Value"
            for j, (label, val) in enumerate(zip(chart_info["x_labels"], chart_info["y_data"])):
                row = table.rows[j + 1].cells
                row[0].text = str(label)
                row[1].text = str(val)
        else:
            buf = chart_buffers[i] if i < len(chart_buffers) else None
            if buf:
                document.add_paragraph(f"Figure {i+1}: {chart_info['chart_title']}")
                document.add_picture(buf, width=Inches(5.5))

    # Recommendations
    document.add_heading("Recommendations", level=1)
    for rec in report_data.get("recommendations", []):
        document.add_paragraph(rec, style="List Number")

    # References
    document.add_heading("References", level=1)
    for cite in report_data.get("citations", []):
        document.add_paragraph(cite, style="List Bullet")

    out = io.BytesIO()
    document.save(out)
    out.seek(0)
    return out.read()


# ===============================
# UI: File Uploader & Run
# ===============================
uploaded_files = st.file_uploader(
    "Upload internal documents (PDF, DOCX, CSV) — you can select multiple files",
    type=["pdf", "docx", "csv", "txt"],
    accept_multiple_files=True,
)

col_run, col_preview = st.columns([1, 3])

with col_run:
    can_run = st.button("🚀 Generate Report", type="primary", use_container_width=True)

with col_preview:
    show_preview = st.checkbox("Show chart previews in app", value=True)

if can_run:
    try:
        if not api_key and not os.environ.get("GEMINI_API_KEY"):
            st.error("Please provide a Gemini API key in the sidebar or via GEMINI_API_KEY.")
            st.stop()

        if not uploaded_files:
            st.error("Please upload at least one file.")
            st.stop()

        with st.spinner("🔎 Parsing & chunking files..."):
            internal_text = load_multiple_files(uploaded_files, chunk_size=chunk_size, overlap=overlap)

        # Combine full context
        full_context = f"INTERNAL DOCUMENTATION:\n{internal_text}\n\nEXTERNAL MARKET DATA:\n{external_text}"

        with st.spinner("🤖 Calling Gemini to generate structured report..."):
            client = build_gemini_client(api_key or os.environ.get("GEMINI_API_KEY"))
            report_data = generate_structured_report(client, model_name, full_context)

        st.success("✅ Structured report generated!")

        # Generate charts
        chart_buffers = []
        if report_data.get("data_for_charts"):
            with st.spinner("📊 Generating charts..."):
                for chart_info in report_data["data_for_charts"]:
                    if chart_info["chart_type_suggestion"].lower() != "table":
                        buf = generate_chart(chart_info)
                        chart_buffers.append(buf)
                    else:
                        chart_buffers.append(None)

        # Optional: Show previews
        if show_preview and report_data.get("data_for_charts"):
            st.subheader("Chart Previews")
            for i, (chart_info, buf) in enumerate(zip(report_data["data_for_charts"], chart_buffers)):
                if chart_info["chart_type_suggestion"].lower() != "table" and buf:
                    st.image(buf, caption=f"Figure {i+1}: {chart_info['chart_title']}", use_container_width=True)
                elif chart_info["chart_type_suggestion"].lower() == "table":
                    df_table = pd.DataFrame(
                        {"Metric": chart_info["x_labels"], "Value": chart_info["y_data"]}
                    )
                    st.markdown(f"**Table {i+1}: {chart_info['chart_title']}**")
                    st.dataframe(df_table, use_container_width=True)

        # Create DOCX bytes and provide download button
        with st.spinner("📝 Building DOCX..."):
            docx_bytes = create_docx_report_bytes(report_data, chart_buffers)

        default_filename = (report_data.get("title", "Automated Report")
                            .replace(" ", "_")
                            .replace(":", "")) + ".docx"

        st.download_button(
            label="⬇️ Download Report (DOCX)",
            data=docx_bytes,
            file_name=default_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        # Quick textual summary in app
        st.subheader("Executive Summary")
        st.write(report_data.get("executive_summary", ""))

        st.subheader("Key Insights")
        for k in report_data.get("key_insights", []):
            st.markdown(f"- {k}")

        st.subheader("Recommendations")
        for r in report_data.get("recommendations", []):
            st.markdown(f"1. {r}")

    except Exception as e:
        st.error(f"❌ Error: {e}")


